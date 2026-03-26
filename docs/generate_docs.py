from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INTUIT_BLUE = RGBColor(0x00, 0x77, 0xC5)
DARK_GRAY   = RGBColor(0x33, 0x33, 0x33)
TEAL        = RGBColor(0x2D, 0xD4, 0xBF)


def set_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        if color:
            run.font.color.rgb = color
        run.font.bold = True
    return h


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "0077C5")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:val"), "clear")
        cell.paragraphs[0]._p.get_or_add_pPr().append(shading)
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        fill = "EBF5FB" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), fill)
            shading.set(qn("w:color"), "auto")
            shading.set(qn("w:val"), "clear")
            cell.paragraphs[0]._p.get_or_add_pPr().append(shading)
    return table


def add_bullet(doc, text, indent=0):
    p = doc.add_paragraph(style="List Bullet" if indent == 0 else "List Bullet 2")
    p.add_run(text)


def spacer(doc):
    doc.add_paragraph("")


# ── DOC 1: PRD ────────────────────────────────────────────────────────────────

def build_prd():
    doc = Document()
    doc.core_properties.title = "MindGym — Product Requirements Document"

    set_heading(doc, "MindGym — Product Requirements Document", 1, INTUIT_BLUE)
    doc.add_paragraph("Version 1.0  ·  March 2026  ·  Status: Draft for Review")
    spacer(doc)

    set_heading(doc, "1. Problem Statement", 2)
    doc.add_paragraph(
        "High-stakes performance situations — job interviews, public speeches, auditions, "
        "presentations — trigger anxiety in the majority of people. Existing tools fall into "
        "two categories: generic meditation apps (Calm, Headspace) that offer no "
        "scenario-specific guidance, and professional sports psychology coaching that is "
        "expensive and inaccessible. There is no self-serve, affordable tool that teaches "
        "structured mental rehearsal for everyday performers."
    )
    p = doc.add_paragraph()
    r = p.add_run("Core insight: ")
    r.bold = True
    p.add_run("Elite athletes have used visualization for decades. This app brings those techniques to anyone facing a high-stakes moment.")
    spacer(doc)

    set_heading(doc, "2. Target Users", 2)
    add_table(doc,
        ["Persona", "Description", "Key Need"],
        [
            ["The Career Climber (Primary)", "Age 25–40, professional. Upcoming interview, pitch, or board presentation. Tried meditation apps, found them generic.", "Scenario-specific mental prep"],
            ["The Performer (Secondary)", "Musician, actor, public speaker, competitive athlete. Performance anxiety is recurring.", "Structured rehearsal, not breathing exercises"],
            ["The Student (Tertiary)", "Age 18–25. Exams, oral defenses, auditions. Price-sensitive.", "Accessible, affordable, quick to use"],
        ]
    )
    spacer(doc)

    set_heading(doc, "3. Core Use Cases", 2)
    add_table(doc,
        ["Use Case", "User Goal"],
        [
            ["Job Interview Prep", "Rehearse confidence, reduce fear of judgment, anchor calm before walking in"],
            ["Public Speaking", "Manage physical anxiety, visualize delivery, build stage presence mentally"],
            ["Sports / Athletic Performance", "Pre-competition focus, simulate execution under pressure, prime muscle memory"],
            ["Audition / Creative Performance", "Reduce perfectionism paralysis, rehearse emotional presence"],
            ["High-Stakes Presentation", "Visualize audience engagement, manage imposter syndrome, anchor authority"],
            ["General Anxiety (catch-all)", "For users who don't fit a specific category but face a stressful moment"],
        ]
    )
    spacer(doc)

    set_heading(doc, "4. Feature List", 2)
    set_heading(doc, "MVP Features", 3)
    mvp_features = [
        "Situation selector — 6 categories",
        "Pre-session personalization: 3 questions (event name, date, #1 goal)",
        "Guided audio sessions: 3-phase format (Calm → Rehearse → Anchor)",
        "8 total sessions across 3 launch categories (Interview ×3, Presentation ×3, Athletic ×2)",
        "Custom audio player: play/pause, scrub, phase indicator",
        "Session completion tracking (list view)",
        "Streak counter",
        "Pre/post mood check-in (1–5 emoji scale)",
        "Anonymous first session — no account required to try",
        "Soft account creation gate post-session",
        "Free tier (3 sessions/month) + Pro tier ($12/month) + Annual ($79/year)",
        "7-day free trial on Pro",
        "Landing page with single CTA",
    ]
    for f in mvp_features:
        add_bullet(doc, f)
    spacer(doc)

    set_heading(doc, "Future Scope (v1.1+)", 3)
    future = [
        "AI-generated personalized scripts (Claude API)",
        "Musical performance and academic exam categories",
        "Mobile apps (iOS, Android)",
        "Email reminder sequences",
        "Mood delta analytics chart",
        "Social sharing and streaks",
        "Wearable integration (HRV, biofeedback)",
        "B2B / team plans (HR, L&D)",
        "Speed controls (0.75×, 1×, 1.25×)",
        "Offline mode / PWA",
    ]
    for f in future:
        add_bullet(doc, f)
    spacer(doc)

    set_heading(doc, "5. User Stories", 2)
    stories = [
        ("Onboarding", [
            "As a first-time visitor, I want to try a session without creating an account so that I can evaluate the product before committing.",
            "As a new user, I want to pick my situation category so that the session is relevant to my actual challenge.",
            "As a user, I want to answer 3 quick personalization questions so that the session feels tailored to me.",
        ]),
        ("Sessions", [
            "As a user, I want to see which phase of the session I'm in so that I know where I am in the journey.",
            "As a user, I want to pause and resume a session so that I don't lose progress if I'm interrupted.",
            "As a user, I want a brief mood check-in before and after the session so that I can track how I feel.",
        ]),
        ("Progress", [
            "As a returning user, I want to see my completed sessions and streak so that I feel motivated to keep going.",
            "As a user, I want to repeat a session I found helpful so that I can reinforce the visualization.",
        ]),
        ("Edge Cases", [
            "As a user who dropped off mid-session, I want the option to restart or resume so that I don't start from scratch.",
            "As a user whose event is tomorrow, I want a shorter session option so that I can prep quickly.",
        ]),
    ]
    for category, items in stories:
        set_heading(doc, category, 3)
        for s in items:
            add_bullet(doc, s)
    spacer(doc)

    set_heading(doc, "6. Success Metrics", 2)
    add_table(doc,
        ["Metric", "MVP Target", "Measurement Method"],
        [
            ["D1 Retention", "60%", "≥1 session completed on signup day"],
            ["D7 Retention", "40%", "Return within 7 days"],
            ["Session Completion Rate", "70%", "User reaches Phase 3 (Anchor)"],
            ["Free → Pro Conversion", "8% within 30 days", "Stripe events"],
            ["Pre/Post Mood Delta", "+1.2 avg", "Self-reported 1–5 scale"],
            ["NPS", "40+", "In-app survey at Day 14"],
            ["Landing Page Bounce Rate", "<55%", "PostHog analytics"],
        ]
    )
    spacer(doc)

    set_heading(doc, "7. Out of Scope (MVP)", 2)
    oos = [
        "Mobile apps (iOS or Android)",
        "AI-generated scripts — template substitution only in MVP",
        "Social / sharing features",
        "Email reminders",
        "Analytics charts for users",
        "B2B / team plans",
        "Live coaching or community",
        "Wearable integrations",
        "Full SEO content strategy",
        "Custom background music selection",
    ]
    for item in oos:
        add_bullet(doc, item)

    doc.save("/Users/clairezhu/Documents/MindGym_Docs/01_PRD.docx")
    print("Saved 01_PRD.docx")


# ── DOC 2: User Flow ──────────────────────────────────────────────────────────

def build_user_flow():
    doc = Document()
    set_heading(doc, "MindGym — User Flow & Experience Design", 1, INTUIT_BLUE)
    doc.add_paragraph("Version 1.0  ·  March 2026  ·  Status: Draft for Review")
    spacer(doc)

    set_heading(doc, "1. End-to-End User Journey", 2)
    steps = [
        ("Landing Page", "User arrives, reads value proposition"),
        ("'Start your free session' CTA", "Entry point — no friction"),
        ("Situation Selector", "User picks from 6 category cards"),
        ("Pre-Session Setup", "3 personalization questions answered"),
        ("Mood Check-In (Pre)", "1–5 emoji scale — baseline captured"),
        ("Session Player", "3-phase guided audio session plays"),
        ("Mood Check-In (Post)", "1–5 emoji scale — delta captured"),
        ("Session Complete Screen", "Mood delta shown + soft account gate"),
        ("Dashboard", "History, streak, upcoming event countdown"),
    ]
    add_table(doc, ["Step", "What Happens"], steps)
    spacer(doc)

    set_heading(doc, "Key Principle", 3)
    p = doc.add_paragraph()
    r = p.add_run("First session plays before account creation. We earn the signup.")
    r.bold = True
    r.font.color.rgb = INTUIT_BLUE
    spacer(doc)

    set_heading(doc, "2. Decision Tree: Routing to the Right Session", 2)
    routing = [
        ("Job Interview", "First round / phone screen", "Session A — anxiety + clarity"),
        ("Job Interview", "Final round / in-person", "Session B — confidence + presence"),
        ("Job Interview", "Negotiation / offer", "Session C — assertiveness + calm"),
        ("Presentation", "Small team / internal", "Session A — clarity + authority"),
        ("Presentation", "Large audience / keynote", "Session B — stage presence + energy"),
        ("Presentation", "Virtual / recorded", "Session C — focus + authenticity"),
        ("Athletic Event", "Individual competition", "Session A — focus + execution"),
        ("Athletic Event", "Team sport", "Session B — composure + trust"),
        ("Audition / Exam", "Any", "v1.1 — deferred"),
        ("General", "Any", "Catch-all session"),
    ]
    add_table(doc, ["Category", "Sub-Type", "Session Routed"], routing)
    spacer(doc)

    set_heading(doc, "MVP Note", 3)
    doc.add_paragraph(
        "In MVP, routing is simplified: the user's 3 personalization answers are injected "
        "into a pre-written script template. Full decision-tree routing with distinct sessions "
        "per sub-type is a v1.1 feature."
    )
    spacer(doc)

    set_heading(doc, "3. Screen-by-Screen Flow", 2)
    screens = [
        ("Landing Page", "Hero headline, 3-step explainer, advisor credibility, pricing, single CTA", "Click CTA → Situation Selector"),
        ("Situation Selector", "6 category cards with icons. 'Most popular' badge on top 2.", "Click card → Pre-Session Setup"),
        ("Pre-Session Setup", "3-question form: event name (text), event date (picker), #1 goal (text)", "Submit → Mood Check-In (pre)"),
        ("Mood Check-In (Pre)", "'How are you feeling right now?' — 5 emoji options. Auto-advances on tap.", "Select emoji → Session Player"),
        ("Session Player", "Dark immersive UI. Breathing animation. Phase indicator. Audio controls. Optional subtitles.", "Complete → Mood Check-In (post). Drop off → saved as incomplete."),
        ("Mood Check-In (Post)", "'How are you feeling now?' — same 5 emoji options.", "Select emoji → Session Complete"),
        ("Session Complete", "Mood delta display. Streak update. Soft account gate CTA.", "Create account → Dashboard. Skip → Situation Selector (once only)."),
        ("Dashboard", "Upcoming event countdown. Session history list. Streak. Start new session CTA. Upgrade prompt.", "Ongoing usage hub"),
    ]
    add_table(doc, ["Screen", "What User Sees & Can Do", "What Happens Next"], screens)
    spacer(doc)

    set_heading(doc, "4. Session Structure", 2)
    add_table(doc,
        ["Phase", "Duration", "What Happens"],
        [
            ["CALM", "1–2 min", "Body scan and breath regulation, tension release, arrive in present"],
            ["REHEARSE", "3–4 min", "Scene-setting, walk through the moment, visualize success + obstacles, emotional rehearsal"],
            ["ANCHOR", "1 min", "Lock in the feeling, physical anchor gesture, return to present with intention"],
            ["TOTAL", "5–6 min", "Typical session length"],
        ]
    )
    spacer(doc)

    set_heading(doc, "5. Edge Cases", 2)
    add_table(doc,
        ["Scenario", "Handling"],
        [
            ["User drops off mid-session", "Mark as incomplete. On next visit show 'Resume your session?' prompt."],
            ["User wants to repeat a session", "'Repeat' button on session history. Re-runs same session, re-does mood check-in."],
            ["User's event is today", "Date picker triggers same-day framing in narration; nudges toward shorter session."],
            ["User skips account creation", "Allow once. Second session completion requires account."],
            ["Audio permissions denied", "Show fallback text-only mode with timed text reveal."],
            ["Audio fails to load", "Show error with retry button; offer text-only fallback."],
        ]
    )

    doc.save("/Users/clairezhu/Documents/MindGym_Docs/02_UserFlow.docx")
    print("Saved 02_UserFlow.docx")


# ── DOC 3: Content Architecture ───────────────────────────────────────────────

def build_content():
    doc = Document()
    set_heading(doc, "MindGym — Content Architecture", 1, INTUIT_BLUE)
    doc.add_paragraph("Version 1.0  ·  March 2026  ·  Status: Draft for Review")
    spacer(doc)

    set_heading(doc, "1. Situation Categories", 2)
    add_table(doc,
        ["Category", "Launch", "Emotions Targeted", "Visualization Focus"],
        [
            ["Job Interview", "✅ MVP", "Anxiety, self-doubt, fear of judgment", "Confidence walking in, articulating clearly, handling tough questions"],
            ["Presentation / Public Speaking", "✅ MVP", "Stage fright, imposter syndrome, physical tension", "Commanding the room, audience engagement, steady presence"],
            ["Athletic Event", "✅ MVP", "Performance anxiety, over-thinking, pressure", "Executing perfectly, composure under pressure, trusting training"],
            ["Audition / Creative", "v1.1", "Perfectionism, vulnerability, creative block", "Authentic expression, letting go of outcome, channeling emotion"],
            ["Exam / Academic", "v1.1", "Overwhelm, blank-mind fear, time pressure", "Clear recall, steady focus, working through difficulty calmly"],
            ["General / High-Stakes", "✅ MVP (catch-all)", "Undefined anxiety, anticipatory stress", "Generic success visualization, calm and grounded presence"],
        ]
    )
    spacer(doc)

    set_heading(doc, "2. Per-Category Script Focus", 2)

    categories = [
        ("Job Interview", [
            ("Opening anchor", "You are prepared. Your experience is real."),
            ("Scene", "Walking into the building, greeting the interviewer warmly, sitting down with ease."),
            ("Core rehearsal", "Answering the hardest question clearly and confidently. The interviewer nodding. You feeling in flow."),
            ("Obstacle handling", "A question you don't know — you pause, breathe, answer thoughtfully. That's strength, not weakness."),
            ("Closing anchor", "You walk out knowing you gave your best. Whatever happens, you showed up fully."),
        ]),
        ("Presentation / Public Speaking", [
            ("Opening anchor", "Your voice carries weight. People want to hear what you have to say."),
            ("Scene", "Walking to the front of the room, scanning the audience, making eye contact."),
            ("Core rehearsal", "Opening line delivered powerfully. The room settling in, listening. You feeling connected to the material."),
            ("Obstacle handling", "An unexpected question from the audience — you welcome it, breathe, respond calmly."),
            ("Closing anchor", "Applause. A feeling of completion. You did it."),
        ]),
        ("Athletic Event", [
            ("Opening anchor", "Your body knows what to do. Your training is in your cells."),
            ("Scene", "Pre-competition environment — the arena, the track, the court. The sounds and smells."),
            ("Core rehearsal", "Executing the key moment in perfect form. Focus without force. Trusting your body."),
            ("Obstacle handling", "A mistake mid-competition — you reset immediately, no dwelling. Next play."),
            ("Closing anchor", "Crossing the finish line. Pride regardless of scoreboard."),
        ]),
    ]

    for cat_name, elements in categories:
        set_heading(doc, cat_name, 3)
        add_table(doc, ["Element", "Content Focus"], elements)
        spacer(doc)

    set_heading(doc, "3. Session Script Template", 2)
    doc.add_paragraph(
        "All sessions follow a 3-phase narrative arc. Template variables are injected "
        "from the user's pre-session answers: [EVENT_NAME], [EVENT_DATE], [GOAL_ACTION]."
    )
    spacer(doc)

    phases = [
        ("CALM Phase", (
            "Close your eyes... take a slow breath in through your nose... hold for a moment... "
            "and release through your mouth. Let your shoulders drop. Let your jaw unclench. "
            "You have nowhere to be right now except here.\n\n"
            "Scan your body from the top of your head to your feet. Wherever you notice tension, "
            "breathe into it and let it soften.\n\nYou are safe. You are ready to prepare."
        )),
        ("REHEARSE Phase", (
            "Now, let's travel to [EVENT_LOCATION / the place where your [EVENT_NAME] will happen]. "
            "Picture it clearly. What do you see? What does it smell like?\n\n"
            "Now see yourself arriving. You are calm. You are grounded.\n\n"
            "Picture the moment you begin. Your [GOAL_ACTION].\n\n"
            "Now, something unexpected happens... Watch how you handle it. "
            "You continue. You are stronger for it."
        )),
        ("ANCHOR Phase", (
            "Take a deep breath and feel what you feel right now — this confidence, this readiness. "
            "Place your hand on your heart, or press your thumb and index finger together. "
            "This is your anchor.\n\n"
            "When you're ready, gently open your eyes. "
            "You are prepared. Go show them what you've got."
        )),
    ]

    for phase_name, script in phases:
        set_heading(doc, phase_name, 3)
        p = doc.add_paragraph(script)
        p.paragraph_format.left_indent = Inches(0.3)
        spacer(doc)

    set_heading(doc, "4. Tone & Voice Guidelines", 2)
    add_table(doc,
        ["Dimension", "Choice", "Rationale"],
        [
            ["Warmth vs. Clinical", "Warm, human", "Users are anxious — clinical tone increases distance"],
            ["Calm vs. Motivational", "Calm-first, motivational at close", "Match user's nervous state at start; build toward energy"],
            ["Pace", "Slow, deliberate", "Mirrors breath pacing; reduces cognitive load"],
            ["Person", "Second person ('you')", "Direct, immersive, personal"],
            ["Vocabulary complexity", "Simple", "Accessible; anxiety impairs working memory"],
            ["Sentence length", "Short during breathing, longer during visualization", "Rhythm mirrors the session arc"],
        ]
    )
    spacer(doc)

    set_heading(doc, "Avoid", 3)
    avoids = [
        "Toxic positivity ('Everything will be perfect!')",
        "Pressure language ('You MUST succeed')",
        "Medical or therapeutic claims",
        "Jargon or complex vocabulary",
    ]
    for a in avoids:
        add_bullet(doc, a)
    spacer(doc)

    set_heading(doc, "5. Static vs. Dynamic Content", 2)
    add_table(doc,
        ["Approach", "MVP?", "Description"],
        [
            ["Static scripts + template substitution", "✅ Yes", "User's 3 answers injected into fixed narrative templates. Simple, fast, no API dependency."],
            ["LLM-generated scripts (Claude API)", "v1.1", "User answers feed a prompt; script generated fresh per session. Deeply personalized. Requires AI API + latency management."],
        ]
    )

    doc.save("/Users/clairezhu/Documents/MindGym_Docs/03_ContentArchitecture.docx")
    print("Saved 03_ContentArchitecture.docx")


# ── DOC 4: Technical Architecture ────────────────────────────────────────────

def build_tech():
    doc = Document()
    set_heading(doc, "MindGym — Technical Architecture", 1, INTUIT_BLUE)
    doc.add_paragraph("Version 1.0  ·  March 2026  ·  Status: Draft for Review")
    spacer(doc)

    set_heading(doc, "1. Recommended Tech Stack", 2)
    add_table(doc,
        ["Layer", "Technology", "Justification"],
        [
            ["Frontend", "Next.js 14 (App Router) + TypeScript", "SSR for fast initial load + SEO; serverless API routes; strong ecosystem"],
            ["Styling", "Tailwind CSS + Framer Motion", "Tailwind for rapid UI; Framer Motion for breathing animations and text reveals"],
            ["Auth", "Supabase Auth", "Email + Google SSO out of the box; integrates with Supabase DB; generous free tier"],
            ["Database", "Supabase (Postgres)", "Row-level security built in; real-time capable; generous free tier"],
            ["Audio", "Custom <audio> wrapper component", "No third-party dependency; phase tracking via React state synced to timeupdate events"],
            ["Payments", "Stripe", "Industry standard; Checkout handles free trial + subscription logic cleanly"],
            ["Hosting", "Vercel", "Zero-config Next.js deployment; edge network; preview deployments per PR"],
            ["Analytics", "PostHog", "Open source; session recording + funnels; self-hostable if privacy required"],
            ["AI (v1.1)", "Anthropic Claude API", "Personalized script generation; streamed response feeds into audio TTS pipeline"],
        ]
    )
    spacer(doc)

    set_heading(doc, "2. Component Breakdown", 2)
    add_table(doc,
        ["Directory", "Component", "Purpose"],
        [
            ["layout/", "Header.tsx", "Nav, auth state, logo"],
            ["layout/", "Footer.tsx", "Minimal links"],
            ["layout/", "PageWrapper.tsx", "Consistent padding / max-width"],
            ["landing/", "HeroSection.tsx", "Headline, subhead, CTA button"],
            ["landing/", "HowItWorks.tsx", "3-step explainer"],
            ["landing/", "PricingSection.tsx", "Free / Pro / Annual cards"],
            ["landing/", "AdvisorCredibility.tsx", "Advisor quote + badge"],
            ["onboarding/", "SituationSelector.tsx", "6 category cards grid"],
            ["onboarding/", "PreSessionForm.tsx", "3-question personalization form"],
            ["onboarding/", "MoodCheckIn.tsx", "1–5 emoji selector (reused pre + post)"],
            ["player/", "SessionPlayer.tsx", "Orchestrates entire session"],
            ["player/", "AudioController.tsx", "Play/pause, scrub, time display"],
            ["player/", "PhaseIndicator.tsx", "CALM / REHEARSE / ANCHOR progress bar"],
            ["player/", "BreathingAnimation.tsx", "SVG circle expand/contract"],
            ["player/", "NarrationSubtitle.tsx", "Optional timed text reveal"],
            ["dashboard/", "SessionHistory.tsx", "List of completed sessions"],
            ["dashboard/", "StreakCounter.tsx", "Flame icon + count"],
            ["dashboard/", "UpcomingEvent.tsx", "Countdown to user's event"],
            ["dashboard/", "UpgradePrompt.tsx", "Free tier upgrade CTA"],
            ["shared/", "Button.tsx", "Variants: primary / ghost / outline"],
            ["shared/", "Modal.tsx", "Account creation gate"],
            ["shared/", "EmojiScale.tsx", "Reusable 1–5 mood selector"],
            ["shared/", "ProgressBar.tsx", "Generic stepped progress"],
        ]
    )
    spacer(doc)

    set_heading(doc, "3. State Management", 2)
    doc.add_paragraph(
        "Approach: React Context + useReducer for session state. Supabase client for server "
        "state. No Redux — overkill for MVP scope."
    )
    spacer(doc)
    set_heading(doc, "SessionContext Shape", 3)
    fields = [
        ("currentCategory", "string", "Selected situation category"),
        ("personalizationAnswers", "object", "Event name, date, goal from pre-session form"),
        ("preMoodScore", "1–5", "Mood captured before session"),
        ("postMoodScore", "1–5", "Mood captured after session"),
        ("currentPhase", "calm | rehearse | anchor | complete", "Current session phase"),
        ("audioProgress", "seconds", "Current playback position"),
        ("isPlaying", "boolean", "Audio playback state"),
        ("sessionId", "uuid", "Assigned at session start"),
    ]
    add_table(doc, ["Field", "Type", "Description"], fields)
    spacer(doc)

    set_heading(doc, "4. Data Model", 2)
    set_heading(doc, "profiles table", 3)
    add_table(doc,
        ["Column", "Type", "Notes"],
        [
            ["id", "uuid", "References Supabase auth.users"],
            ["stripe_customer_id", "text", "Stripe customer reference"],
            ["subscription_tier", "text", "'free' | 'pro'"],
            ["trial_ends_at", "timestamptz", "Null if not on trial"],
        ]
    )
    spacer(doc)
    set_heading(doc, "sessions table", 3)
    add_table(doc,
        ["Column", "Type", "Notes"],
        [
            ["id", "uuid", "Primary key"],
            ["user_id", "uuid", "Nullable — anonymous first session allowed"],
            ["category", "text", "e.g. 'interview', 'presentation'"],
            ["session_key", "text", "e.g. 'interview_final_round'"],
            ["event_name", "text", "From pre-session question 1"],
            ["event_date", "date", "From pre-session question 2"],
            ["user_goal", "text", "From pre-session question 3"],
            ["pre_mood", "int", "1–5, captured before session"],
            ["post_mood", "int", "1–5, null until session complete"],
            ["started_at", "timestamptz", "Session start timestamp"],
            ["completed_at", "timestamptz", "Null if dropped off"],
            ["phase_reached", "text", "calm | rehearse | anchor | complete"],
        ]
    )
    spacer(doc)

    set_heading(doc, "5. Third-Party Integrations", 2)
    add_table(doc,
        ["Integration", "Purpose", "MVP?"],
        [
            ["Supabase Auth", "User authentication (email + Google SSO)", "✅ Yes"],
            ["Stripe", "Subscriptions + free trial logic", "✅ Yes"],
            ["PostHog", "Analytics + funnel tracking + session recording", "✅ Yes"],
            ["Audio CDN (Cloudflare R2)", "Serve licensed audio files at low latency", "✅ Yes"],
            ["Claude API (Anthropic)", "AI-generated personalized scripts", "v1.1"],
            ["Resend / Postmark", "Email reminders and transactional email", "v1.1"],
        ]
    )
    spacer(doc)

    set_heading(doc, "6. Performance Considerations", 2)
    perf = [
        "Audio preloading: Preload next-phase audio segment while current phase plays to eliminate gaps between phases.",
        "Breathing animation: SVG-based (not canvas) — GPU-composited, no layout thrash.",
        "LCP target <2.5s: Hero section is static HTML + CSS; no blocking JS. Audio and animations are lazy-loaded.",
        "Mobile: All interactions designed touch-first. Audio player controls are 48px tap targets minimum.",
        "Background playback: Standard <audio> with playsinline — tab switch must not stop audio.",
        "No autoplay: Browser policy + intentionality — user explicitly presses Play.",
    ]
    for p_item in perf:
        add_bullet(doc, p_item)

    doc.save("/Users/clairezhu/Documents/MindGym_Docs/04_TechnicalArchitecture.docx")
    print("Saved 04_TechnicalArchitecture.docx")


# ── DOC 5: Design System ──────────────────────────────────────────────────────

def build_design():
    doc = Document()
    set_heading(doc, "MindGym — Design System Proposal", 1, INTUIT_BLUE)
    doc.add_paragraph("Version 1.0  ·  March 2026  ·  Status: Draft for Review")
    spacer(doc)

    set_heading(doc, "1. Color Palette", 2)
    doc.add_paragraph(
        "Principle: The palette should feel like 2am by the ocean — dark, calm, spacious. "
        "Avoid clinical whites or high-contrast brights during sessions."
    )
    spacer(doc)
    add_table(doc,
        ["Role", "Name", "Hex", "Usage"],
        [
            ["Primary Background", "Deep Ocean", "#0A1628", "Page backgrounds, dark mode base"],
            ["Secondary Background", "Midnight Blue", "#0D2137", "Card backgrounds, player UI"],
            ["Accent (primary)", "Calm Teal", "#2DD4BF", "CTAs, phase indicators, progress bars"],
            ["Accent (warm)", "Soft Gold", "#F59E0B", "Streak counter, positive mood states"],
            ["Text Primary", "Off-White", "#F0F4F8", "Body text on dark backgrounds"],
            ["Text Muted", "Slate", "#94A3B8", "Captions, secondary labels"],
            ["Success", "Sage Green", "#4ADE80", "Completion states, positive mood"],
            ["Neutral / Warning", "Warm Sand", "#FCD34D", "Mid-mood states"],
            ["Error / Tension", "Muted Rose", "#F87171", "Low mood states, alerts"],
        ]
    )
    spacer(doc)

    set_heading(doc, "2. Typography", 2)
    add_table(doc,
        ["Role", "Font", "Weight", "Size"],
        [
            ["Display / Hero", "Inter or DM Sans", "700 Bold", "48–64px"],
            ["Section Headings", "Inter", "600 SemiBold", "28–36px"],
            ["Body", "Inter", "400 Regular", "16–18px"],
            ["Session Narration", "Lora (serif)", "400 Regular Italic", "20–24px"],
            ["Captions / Labels", "Inter", "400 Regular", "12–14px"],
        ]
    )
    spacer(doc)
    p = doc.add_paragraph()
    r = p.add_run("Note: ")
    r.bold = True
    p.add_run(
        "Narration text uses serif italic (Lora) to feel distinct from UI text — "
        "signaling to the user that they have entered 'the experience.'"
    )
    spacer(doc)

    set_heading(doc, "3. Animation & Transition Style", 2)
    add_table(doc,
        ["Element", "Animation", "Duration", "Easing"],
        [
            ["Breathing circle (CALM phase)", "Scale 1.0 → 1.35 → 1.0 loop", "4s in / 4s out", "ease-in-out"],
            ["Phase transition", "Crossfade background + label", "800ms", "ease-in-out"],
            ["Narration text reveal", "Fade in line-by-line", "300ms per line", "ease-out"],
            ["Page transitions", "Fade + subtle upward translate", "400ms", "ease-out"],
            ["Mood emoji selection", "Scale bounce on tap", "150ms", "Spring"],
            ["Session complete", "Soft radial pulse from center", "1.5s, 2 pulses", "ease-out"],
        ]
    )
    spacer(doc)
    p = doc.add_paragraph()
    r = p.add_run("Principle: ")
    r.bold = True
    p.add_run(
        "Animations serve the emotional arc — they should slow the user down, not entertain them. "
        "Nothing flashy during sessions."
    )
    spacer(doc)

    set_heading(doc, "4. Key UI Patterns", 2)

    patterns = [
        ("Phase Indicator",
         "Three segments (CALM / REHEARSE / ANCHOR) as a horizontal bar at the top of the player. "
         "Active segment glows in Calm Teal. Completed segments solid. Inactive segments muted."),
        ("Breathing Animation",
         "Centered SVG circle with soft gradient fill (teal to deep blue). Expands on inhale cue, "
         "contracts on exhale. Optional text label ('breathe in... breathe out...') fades beneath."),
        ("Narration Subtitles",
         "Optional overlay at bottom of player. Lines fade in one at a time, synced to audio timestamps. "
         "Max 2 lines visible at once. Muted and dismissible — accessibility feature, not primary UX."),
        ("Emoji Mood Scale",
         "5 large emoji rendered as buttons. Selected state: scale up + glow ring in accent color. "
         "No number labels — purely a felt response."),
        ("Soft Account Gate",
         "Post-session modal. Dark overlay, centered card. Headline: 'Save your progress.' "
         "Email input + 'Continue with Google' button. Small 'maybe later' link (allowed once). "
         "Never blocks the session itself."),
    ]

    for name, desc in patterns:
        set_heading(doc, name, 3)
        doc.add_paragraph(desc)
        spacer(doc)

    set_heading(doc, "5. Accessibility Considerations", 2)
    add_table(doc,
        ["Consideration", "Implementation"],
        [
            ["Audio-only content", "All sessions have optional narration subtitles (timed text)"],
            ["Color contrast", "All text meets WCAG AA (4.5:1) on dark backgrounds"],
            ["Keyboard navigation", "All interactive elements focusable; player: spacebar = play/pause, arrow keys = scrub"],
            ["Reduced motion", "Breathing animation and transitions respect prefers-reduced-motion — replaced with static states"],
            ["Screen readers", "Phase changes announced via aria-live region; player controls have aria-label"],
            ["Touch targets", "Minimum 48×48px on all tappable elements"],
            ["Font scaling", "Layout uses rem units; tested up to 200% browser font size"],
        ]
    )

    doc.save("/Users/clairezhu/Documents/MindGym_Docs/05_DesignSystem.docx")
    print("Saved 05_DesignSystem.docx")


if __name__ == "__main__":
    build_prd()
    build_user_flow()
    build_content()
    build_tech()
    build_design()
    print("\nAll 5 documents saved to /Users/clairezhu/Documents/MindGym_Docs/")
